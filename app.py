from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
import sqlite3
import os
import json
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document

app = Flask(__name__, template_folder='Frontend/templates', static_folder='Frontend/static')
app.secret_key = 'tu_clave_secreta'

def get_db_connection():
    db_dir = os.path.join(os.path.dirname(__file__), 'Database')
    db_path = os.path.join(db_dir, 'usuarios.db')
    if not os.path.exists(db_dir):
        os.makedirs(db_dir)
    if not os.path.exists(db_path):
        # Create the database and table if they don't exist
        conn = sqlite3.connect(db_path)
        conn.execute('''CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            apellido_paterno TEXT NOT NULL,
            apellido_materno TEXT NOT NULL,
            correo TEXT NOT NULL UNIQUE,
            contrasena TEXT NOT NULL
        )''')
        # Crear tablas nuevas si no existen
        conn.execute('''CREATE TABLE IF NOT EXISTS docente (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT UNIQUE NOT NULL
        )''')
        conn.execute('''CREATE TABLE IF NOT EXISTS coordinador (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT UNIQUE NOT NULL
        )''')
        conn.execute('''CREATE TABLE IF NOT EXISTS grupo (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT UNIQUE NOT NULL
        )''')
        conn.commit()
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/', methods=['GET'])
def index():
    return render_template('login_registro.html')

@app.route('/registro', methods=['POST'])
def registro():
    nombre = request.form['nombre']
    apellido_paterno = request.form['apellido_pat']
    apellido_materno = request.form['apellido_mat']
    correo = request.form['correo_reg']
    contrasena = request.form['contrasena_reg']
    confirmar = request.form['confirmar_contrasena']

    # Validación de correo institucional
    if not correo.endswith('@uptecamac.edu.mx'):
        flash('Solo se permiten correos institucionales @uptecamac.edu.mx')
        return redirect(url_for('index'))

    # Validación de contraseñas iguales
    if contrasena != confirmar:
        flash('Las contraseñas no coinciden')
        return redirect(url_for('index'))

    hashed_password = generate_password_hash(contrasena)

    conn = get_db_connection()
    try:
        conn.execute(
            'INSERT INTO usuarios (nombre, apellido_paterno, apellido_materno, correo, contrasena) VALUES (?, ?, ?, ?, ?)',
            (nombre, apellido_paterno, apellido_materno, correo, hashed_password)
        )
        conn.commit()
        flash('Usuario registrado exitosamente')
    except sqlite3.IntegrityError:
        flash('El correo ya está registrado')
    finally:
        conn.close()
    return redirect(url_for('index'))

@app.route('/dashboard')
def dashboard():
    nombre = session.get('nombre', 'Usuario')
    return render_template('dashboard.html', nombre=nombre)

@app.route('/login', methods=['POST'])
def login():
    correo = request.form['correo']
    contrasena = request.form['contrasena']

    conn = get_db_connection()
    user = conn.execute(
        'SELECT * FROM usuarios WHERE correo = ?',
        (correo,)
    ).fetchone()
    conn.close()

    if user and check_password_hash(user['contrasena'], contrasena):
        session['nombre'] = user['nombre']
        session['usuario_id'] = user['id']
        flash('Inicio de sesión exitoso')
        return redirect(url_for('dashboard'))
    else:
        flash('Correo o contraseña incorrectos')
        return redirect(url_for('index'))

@app.route('/registrar_evento')
def registrar_evento():
    return render_template('registrar_eventos.html')

# ------------------- CRUD ENDPOINTS PARA DOCENTE, COORDINADOR Y GRUPO -------------------

# --- DOCENTES ---
@app.route('/api/docentes', methods=['GET'])
def get_docentes():
    conn = get_db_connection()
    cur = conn.execute('SELECT id, nombre FROM docente ORDER BY nombre')
    docentes = [dict(row) for row in cur.fetchall()]
    conn.close()
    return jsonify(docentes)

@app.route('/api/docentes', methods=['POST'])
def add_docente():
    nombre = request.json.get('nombre', '').strip()
    if not nombre:
        return jsonify({'error': 'Nombre requerido'}), 400
    conn = get_db_connection()
    try:
        conn.execute('INSERT INTO docente (nombre) VALUES (?)', (nombre,))
        conn.commit()
        return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Docente ya existe'}), 400
    finally:
        conn.close()

@app.route('/api/docentes/<int:docente_id>', methods=['DELETE'])
def delete_docente(docente_id):
    conn = get_db_connection()
    conn.execute('DELETE FROM docente WHERE id = ?', (docente_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# --- COORDINADORES ---
@app.route('/api/coordinadores', methods=['GET'])
def get_coordinadores():
    conn = get_db_connection()
    cur = conn.execute('SELECT id, nombre FROM coordinador ORDER BY nombre')
    coordinadores = [dict(row) for row in cur.fetchall()]
    conn.close()
    return jsonify(coordinadores)

@app.route('/api/coordinadores', methods=['POST'])
def add_coordinador():
    nombre = request.json.get('nombre', '').strip()
    if not nombre:
        return jsonify({'error': 'Nombre requerido'}), 400
    conn = get_db_connection()
    try:
        conn.execute('INSERT INTO coordinador (nombre) VALUES (?)', (nombre,))
        conn.commit()
        return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Coordinador ya existe'}), 400
    finally:
        conn.close()

@app.route('/api/coordinadores/<int:coordinador_id>', methods=['DELETE'])
def delete_coordinador(coordinador_id):
    conn = get_db_connection()
    conn.execute('DELETE FROM coordinador WHERE id = ?', (coordinador_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

# --- GRUPOS ---
@app.route('/api/grupos', methods=['GET'])
def get_grupos():
    conn = get_db_connection()
    cur = conn.execute('SELECT id, nombre FROM grupo ORDER BY nombre')
    grupos = [dict(row) for row in cur.fetchall()]
    conn.close()
    return jsonify(grupos)

@app.route('/api/grupos', methods=['POST'])
def add_grupo():
    nombre = request.json.get('nombre', '').strip()
    if not nombre:
        return jsonify({'error': 'Nombre requerido'}), 400
    conn = get_db_connection()
    try:
        conn.execute('INSERT INTO grupo (nombre) VALUES (?)', (nombre,))
        conn.commit()
        return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Grupo ya existe'}), 400
    finally:
        conn.close()

@app.route('/api/grupos/<int:grupo_id>', methods=['DELETE'])
def delete_grupo(grupo_id):
    conn = get_db_connection()
    conn.execute('DELETE FROM grupo WHERE id = ?', (grupo_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/eventos/borrador', methods=['POST'])
def guardar_borrador():
    datos = request.json.get('datos')
    usuario_id = session.get('usuario_id')  # Ajusta según tu sistema de login
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401
    if not datos.get('fechas') or not isinstance(datos['fechas'], list) or len(datos['fechas']) == 0:
        return jsonify({'error': 'Debes ingresar al menos una fecha.'}), 400
    if not datos.get('horas') or not isinstance(datos['horas'], list) or len(datos['horas']) == 0:
        return jsonify({'error': 'Debes ingresar al menos una hora.'}), 400

    conn = get_db_connection()
    # Busca si ya hay un borrador para este usuario
    borrador = conn.execute(
        'SELECT id FROM evento WHERE usuario_id = ? AND estado = ?', (usuario_id, 'borrador')
    ).fetchone()
    if borrador:
        # Actualiza el borrador existente
        conn.execute(
            'UPDATE evento SET datos = ? WHERE id = ?', (json.dumps(datos), borrador['id'])
        )
    else:
        # Inserta nuevo borrador
        conn.execute(
            'INSERT INTO evento (usuario_id, datos, estado) VALUES (?, ?, ?)',
            (usuario_id, json.dumps(datos), 'borrador')
        )
    conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/eventos/borrador', methods=['GET'])
def obtener_borradores():
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401
    conn = get_db_connection()
    borradores = conn.execute(
        "SELECT id, datos FROM evento WHERE usuario_id = ? AND estado = 'borrador'",
        (usuario_id,)
    ).fetchall()
    conn.close()
    # Extraer el título del evento desde los datos JSON
    eventos = []
    for row in borradores:
        datos = json.loads(row['datos'])
        eventos.append({
            'id': row['id'],
            'titulo': datos.get('titulo', 'Sin título')
        })
    return jsonify(eventos)

@app.route('/evento/<int:evento_id>')
def ver_evento(evento_id):
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return redirect(url_for('index'))
    conn = get_db_connection()
    evento = conn.execute(
        "SELECT * FROM evento WHERE id = ? AND usuario_id = ?", (evento_id, usuario_id)
    ).fetchone()
    conn.close()
    if not evento:
        flash('Evento no encontrado')
        return redirect(url_for('dashboard'))
    datos = json.loads(evento['datos'])
    return render_template('ver_evento.html', evento=evento, datos=datos)

@app.route('/api/eventos/<int:evento_id>', methods=['GET'])
def obtener_evento(evento_id):
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401
    conn = get_db_connection()
    evento = conn.execute(
        "SELECT * FROM evento WHERE id = ? AND usuario_id = ?", (evento_id, usuario_id)
    ).fetchone()
    conn.close()
    if not evento:
        return jsonify({'error': 'Evento no encontrado'}), 404
    datos = json.loads(evento['datos'])
    return jsonify({'datos': datos, 'estado': evento['estado']})

@app.route('/api/eventos/usuario', methods=['GET'])
def obtener_eventos_usuario():
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401
    conn = get_db_connection()
    eventos = conn.execute(
        "SELECT id, datos, estado FROM evento WHERE usuario_id = ?", (usuario_id,)
    ).fetchall()
    conn.close()
    resultado = []
    for row in eventos:
        datos = json.loads(row['datos'])
        resultado.append({
            'id': row['id'],
            'titulo': datos.get('titulo', 'Sin título'),
            'estado': row['estado']
        })
    return jsonify(resultado)


@app.route('/api/eventos/aprobar', methods=['POST'])
def aprobar_evento():
    datos = request.json.get('datos')
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401

    # 1. Validaciones de campos obligatorios
    campos_obligatorios = ['titulo', 'tematica', 'justificacion', 'objetivo', 'dinamica', 'lugar', 'docente', 'coordinador']
    for campo in campos_obligatorios:
        if not datos.get(campo) or not str(datos.get(campo)).strip():
            return jsonify({'error': f'El campo {campo} es obligatorio.'}), 400
    if not datos.get('grupos_asignados') or not isinstance(datos['grupos_asignados'], list) or len(datos['grupos_asignados']) == 0:
        return jsonify({'error': 'Selecciona al menos un grupo.'}), 400
    if not datos.get('recursos_solicitados') or not isinstance(datos['recursos_solicitados'], list) or len(datos['recursos_solicitados']) == 0:
        return jsonify({'error': 'Agrega al menos un recurso solicitado.'}), 400

    # 2. Validaciones de formato y lógica
    if len(datos['justificacion'].strip()) < 20:
        return jsonify({'error': 'La justificación debe tener al menos 20 caracteres.'}), 400
    if len(datos['objetivo'].strip()) < 10:
        return jsonify({'error': 'El objetivo debe tener al menos 10 caracteres.'}), 400

    # Fechas: al menos una válida y futura
    from datetime import datetime
    alguna_fecha_valida = False
    for key, value in datos.items():
        if key.startswith('fecha-dia-') and value:
            try:
                fecha = datetime.fromisoformat(value)
                if fecha > datetime.now():
                    alguna_fecha_valida = True
            except Exception:
                continue
    if not alguna_fecha_valida:
        return jsonify({'error': 'Debes ingresar al menos una fecha futura.'}), 400

    # 3. Validación de unicidad (no permitir eventos con el mismo título para el mismo usuario en estado pendiente/aprobado)
    conn = get_db_connection()
    existe = conn.execute(
        "SELECT id FROM evento WHERE usuario_id = ? AND LOWER(json_extract(datos, '$.titulo')) = ? AND estado IN ('aprobado', 'pendiente')",
        (usuario_id, datos['titulo'].strip().lower())
    ).fetchone()
    if existe:
        conn.close()
        return jsonify({'error': 'Ya existe un evento con ese título en estado pendiente o aprobado.'}), 400

    # 4. Guardado/actualización
    borrador = conn.execute(
        'SELECT id FROM evento WHERE usuario_id = ? AND estado = ?', (usuario_id, 'borrador')
    ).fetchone()
    if borrador:
        conn.execute(
            'UPDATE evento SET datos = ?, estado = ? WHERE id = ?', (json.dumps(datos), 'aprobado', borrador['id'])
        )
    else:
        conn.execute(
            'INSERT INTO evento (usuario_id, datos, estado) VALUES (?, ?, ?)',
            (usuario_id, json.dumps(datos), 'aprobado')
        )
    conn.commit()
    conn.close()
    return jsonify({'success': True})
    datos = request.json.get('datos')
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401

    conn = get_db_connection()
    # Busca si ya hay un borrador para este usuario
    borrador = conn.execute(
        'SELECT id FROM evento WHERE usuario_id = ? AND estado = ?', (usuario_id, 'borrador')
    ).fetchone()
    if borrador:
        # Actualiza el borrador existente a aprobado
        conn.execute(
            'UPDATE evento SET datos = ?, estado = ? WHERE id = ?', (json.dumps(datos), 'aprobado', borrador['id'])
        )
    else:
        # Inserta nuevo evento aprobado
        conn.execute(
            'INSERT INTO evento (usuario_id, datos, estado) VALUES (?, ?, ?)',
            (usuario_id, json.dumps(datos), 'aprobado')
        )
    conn.commit()
    conn.close()
    return jsonify({'success': True})



@app.route('/api/eventos/<int:evento_id>/a_borrador', methods=['POST'])
def evento_a_borrador(evento_id):
    usuario_id = session.get('usuario_id')
    if not usuario_id:
        return jsonify({'error': 'No autenticado'}), 401
    conn = get_db_connection()
    evento = conn.execute(
        "SELECT * FROM evento WHERE id = ? AND usuario_id = ?", (evento_id, usuario_id)
    ).fetchone()
    if not evento:
        conn.close()
        return jsonify({'error': 'Evento no encontrado'}), 404
    # Cambia el estado a 'borrador'
    conn.execute(
        "UPDATE evento SET estado = ? WHERE id = ?", ('borrador', evento_id)
    )
    conn.commit()
    conn.close()
    return jsonify({'success': True})

from docx import Document

CAMPOS_EVENTO = {
    "título": "titulo",
    "titulo": "titulo",
    "categoría": "categoria",
    "impacto": "impacto",
    "temática": "tematica",
    "justificación": "justificacion",
    "objetivo": "objetivo",
    "dinámica": "dinamica",
    "lugar del evento": "lugar",
    "duración del evento": "duracion",
    "periodo cuatrimestral": "periodo_cuatrimestral",
    "recursos solicitados": "recursos",
    "docente": "docente",
    "grupos asignados": "grupos",
    "coordinador": "coordinador"
}

def convertir_a_template(path_docx):
    doc = Document(path_docx)
    for p in doc.paragraphs:
        texto = p.text

        # 1. Lugar y fecha actual
        if "Tecámac, Estado de México" in texto and "a" in texto:
            p.text = "{{evento.lugar_actual}} a {{fecha_actual}}"

        # 2. Oficio No.
        elif "Oficio No." in texto:
            p.text = "Oficio No. {{numero_oficio}}"

        # 3. Categoría
        elif "CATEGORIA:" in texto:
            p.text = "CATEGORIA: {{evento.categoria}}"

        # 4. Evento
        elif "EVENTO:" in texto:
            p.text = "EVENTO: {{evento.titulo}}"

        # 5. Impacto
        elif "Impacto:" in texto and "SEAES:" in texto:
            p.text = "Impacto: {{evento.impacto}}"

        # 6. Temática
        elif "Temática:" in texto:
            p.text = "Tematica: {{evento.tematica}}"

        # 7. Justificación
        elif "Justificación:" in texto:
            p.text = "Justificacion: {{evento.justificacion}}"

        # 8. Objetivo
        elif "Objetivo:" in texto:
            p.text = "Objetivo: {{evento.objetivo}}"

        # 9. Dinámica
        elif "Dinámica:" in texto:
            p.text = "Dinamica: {{evento.dinamica}}"

        # 10. Fecha (puede haber varias)
        elif texto.strip().startswith("Fecha:"):
            p.text = "Fecha: {{evento.fecha}}"

        # 11. Fecha(s)
        elif texto.strip().startswith("Fecha:"):
            p.text = "Fechas:"
            idx = doc.paragraphs.index(p)
            doc.paragraphs.insert(idx + 1, doc.add_paragraph("{% for fecha in evento.fechas %}{{ fecha }}{% endfor %}"))
        # 12. Horario
        elif "Horario:" in texto:
            p.text = "Horarios:"
            idx = doc.paragraphs.index(p)
            doc.paragraphs.insert(idx + 1, doc.add_paragraph("{% for hora in evento.horas %}{{ hora }}{% endfor %}"))

        # 13. Recursos solicitados (lista)
        # Dentro de convertir_a_template
        elif "Recursos solicitados:" in texto:
            p.text = "Recursos solicitados:"
            # Inserta la lista justo después
            idx = doc.paragraphs.index(p)
            doc.paragraphs.insert(idx + 1, doc.add_paragraph("{% for recurso in evento.recursos_solicitados %}• {{ recurso }}{% endfor %}"))
    # Bloque para múltiples eventos
    doc.paragraphs.insert(0, doc.add_paragraph("{% for evento in eventos %}"))
    doc.add_paragraph("{% endfor %}")

    ruta_nueva = path_docx.replace("subidas", "editadas").replace(".docx", "_template.docx")
    doc.save(ruta_nueva)
    return ruta_nueva
    
@app.route("/subir-plantilla", methods=["GET", "POST"])
def subir_plantilla():
    if request.method == "POST":
        archivo = request.files["archivo"]
        if archivo.filename.endswith(".docx"):
            ruta_guardado = os.path.join("templates_plantillas_subidas", archivo.filename)
            archivo.save(ruta_guardado)
            plantilla_convertida = convertir_a_template(ruta_guardado)
            
            # Guarda la metadata en la BD
            guardar_metadata(archivo.filename, plantilla_convertida)
            
            # Leer el contenido para vista previa
            doc = Document(plantilla_convertida)
            contenido = [p.text for p in doc.paragraphs if p.text.strip()]
            # Renderiza la vista previa
            return render_template("vista_previa_plantilla.html", contenido=contenido, plantilla_path=plantilla_convertida)
        else:
            return "Formato no válido. Sube un archivo .docx"
    return render_template("subir_plantilla.html")

@app.route("/descargar_plantilla_editada", methods=["POST"])
def descargar_plantilla_editada():
    plantilla_path = request.form.get("plantilla_path")
    # Aquí podrías aplicar los cambios de los campos editados si lo deseas
    # Por ahora solo envía el archivo original convertido
    return send_file(plantilla_path, as_attachment=True)

def guardar_metadata(nombre, archivo):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO plantillas (nombre, archivo, fecha_creacion) VALUES (?, ?, datetime('now'))", (nombre, archivo))
    conn.commit()
    conn.close()

if __name__ == '__main__':
    app.run(debug=True)
